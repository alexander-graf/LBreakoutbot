
from docx import Document
from docx.shared import Inches
document = Document()
document.add_heading('Document Title', 0)

from docx.enum.style import WD_STYLE_TYPE


p = document.add_paragraph('A plain paragraph having some ')
p.add_run("Fick die Oma in den fetten Arsch")


p.add_run('bold').bold = True
p.add_run(' and some ')
p.add_run('italic.').italic = True
document.add_heading('Heading, level 1', level=1)
document.add_paragraph('Intense quote', style='Intense Quote')

table = document.add_table(rows=1, cols=3)
hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'Qty'
hdr_cells[1].text = 'Id'
hdr_cells[2].text = 'Desc'



paragraph = document.add_paragraph()
paragraph.add_run('Lorem ipsum ')
paragraph.add_run('dolor').bold = True
paragraph.add_run(' sit amet.')

paragraph_format = paragraph.paragraph_format
tab_stops = paragraph_format.tab_stops
tab_stop = tab_stops.add_tab_stop(Inches(2.5))


tab_stop = tab_stops.add_tab_stop(Inches(1.5), WD_TAB_ALIGNMENT.RIGHT, WD_TAB_LEADER.DOTS)


document.save('demo.docx')
