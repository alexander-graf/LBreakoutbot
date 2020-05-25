
from docx import Document
from docx.shared import Inches


from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_TAB_ALIGNMENT, WD_TAB_LEADER





#Definition von tab_stops
tab_stops = paragraph_format.tab_stops

document = Document()
document.add_heading('Document Title', 0)


tab_stop = tab_stops.add_tab_stop(Inches(3.5), WD_TAB_ALIGNMENT.RIGHT, WD_TAB_LEADER.DOTS)


p = document.add_paragraph('A plain paragraph having some ')
p.add_run("Fick die Oma in den fetten Arsch")












document.save('temo.docx')
