#!/usr/bin/env python


import docx    #Holt das Modul für die Bearbeitung von docx Dokumenten
import shared

#Neues Dokument erstellen
#
#Wie mache ich Tabsprünge?
#Dezimal Tabs?
#


doc = docx.Document()
doc.add_paragraph('Hello world!')

doc.add_paragraph("Ja, fick die Waldfee!")

doc_format.right_indent






doc.save('helloworld.docx')   #speichert es ab

