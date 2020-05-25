
from docx import Document
from docx.shared import Inches


from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_TAB_ALIGNMENT, WD_TAB_LEADER
from docx.enum.text import WD_ALIGN_PARAGRAPH


Datum = "11.11.2040"





##Word-Dokument bauen

document = Document()

#Löschen der ersten Zeile
document._body.clear_content()

paragraph = document.add_paragraph("Graf-Computer-Service\t" + Datum)


paragraph_format = paragraph.paragraph_format

#paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
#paragraph_format.alignment

tab_stops = paragraph_format.tab_stops
#tab_stop = tab_stops.add_tab_stop(Inches(1.5))

#rechter TAB
tab_stop = tab_stops.add_tab_stop(Inches(6), WD_TAB_ALIGNMENT.RIGHT, WD_TAB_LEADER.DOTS)

paragraph = document.add_paragraph("Bartholomäusstr. 66")
paragraph = document.add_paragraph("90489 Nürnberg")
paragraph = document.add_paragraph("Telefon: 0179 / 1577763")



document.save('gemo.docx')
